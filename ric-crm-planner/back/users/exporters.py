from __future__ import annotations

from collections import Counter
from datetime import date, datetime
from io import BytesIO
from urllib.parse import quote
import re

from django.http import HttpResponse
from django.utils import timezone

from users.models import Application, Direction, Event, Project

SYSTEM_FIELD_IDS = {"studentName", "telegram", "university", "course", "specialization"}
DEFAULT_FIELD_LABELS = {
    "studentName": "ФИО",
    "telegram": "Аккаунт VK",
    "university": "Университет",
    "course": "Курс",
    "specialization": "Специализация",
    "about": "О себе",
}
FIELD_TYPE_LABELS = {
    "text": "Короткий текст",
    "textarea": "Длинный текст",
    "select": "Выпадающий список",
}


def _stringify(value) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Да" if value else "Нет"
    if isinstance(value, (datetime, date)):
        return _format_datetime(value) if isinstance(value, datetime) else _format_date(value)
    if isinstance(value, (list, tuple, set)):
        return ", ".join(_stringify(item) for item in value if _stringify(item))
    if isinstance(value, dict):
        return "; ".join(f"{key}: {_stringify(item)}" for key, item in value.items())
    return str(value)


def _format_date(value) -> str:
    if not value:
        return ""
    if isinstance(value, datetime):
        value = timezone.localtime(value) if timezone.is_aware(value) else value
        return value.strftime("%d.%m.%Y")
    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")
    return str(value)


def _format_datetime(value) -> str:
    if not value:
        return ""
    if isinstance(value, datetime):
        value = timezone.localtime(value) if timezone.is_aware(value) else value
        return value.strftime("%d.%m.%Y %H:%M")
    return str(value)


def _display_user(user) -> str:
    if not user:
        return ""
    profile = getattr(user, "crm_profile", None)
    if profile:
        parts = [profile.surname, profile.name, profile.patronymic]
        name = " ".join(part for part in parts if part).strip()
        if name:
            return name
    full_name = user.get_full_name().strip() if hasattr(user, "get_full_name") else ""
    return full_name or getattr(user, "email", "") or getattr(user, "username", "") or f"#{user.pk}"


def _profile_value(user, attr: str) -> str:
    profile = getattr(user, "crm_profile", None)
    return _stringify(getattr(profile, attr, "")) if profile else ""


def _event_specializations(event: Event) -> list[str]:
    values = [item.specialization.name for item in event.event_specializations.select_related("specialization").all()]
    if values:
        return values
    return [event.specialization.name] if event.specialization_id else []


def _normalize_form_fields(raw_fields) -> list[dict[str, object]]:
    fields = raw_fields if isinstance(raw_fields, list) else []
    normalized: list[dict[str, object]] = []
    seen: set[str] = set()

    for field in fields:
        if not isinstance(field, dict):
            continue
        field_id = str(field.get("id") or "").strip()
        if not field_id or field_id in seen:
            continue
        seen.add(field_id)
        field_type = field.get("type") if field.get("type") in {"text", "textarea", "select"} else "text"
        label = str(field.get("label") or DEFAULT_FIELD_LABELS.get(field_id) or field_id).strip()
        options = field.get("options") if isinstance(field.get("options"), list) else []
        normalized.append(
            {
                "id": field_id,
                "label": label,
                "type": field_type,
                "required": bool(field.get("required")),
                "options": [str(option).strip() for option in options if str(option).strip()],
            }
        )

    for field_id, label in DEFAULT_FIELD_LABELS.items():
        if field_id not in seen:
            normalized.insert(
                0 if field_id == "studentName" else len([item for item in normalized if item.get("id") in DEFAULT_FIELD_LABELS]),
                {
                    "id": field_id,
                    "label": label,
                    "type": "textarea" if field_id == "about" else "text",
                    "required": field_id in SYSTEM_FIELD_IDS,
                    "options": [],
                },
            )
            seen.add(field_id)

    order = {field_id: index for index, field_id in enumerate(DEFAULT_FIELD_LABELS.keys())}
    return sorted(normalized, key=lambda item: (order.get(str(item["id"]), 1000), str(item["label"])))


def _custom_export_fields(event: Event, applications) -> list[dict[str, object]]:
    fields = [field for field in _normalize_form_fields(event.application_form_fields) if field["id"] not in SYSTEM_FIELD_IDS]
    known = {str(field["id"]) for field in fields}
    extra_keys: set[str] = set()

    for application in applications:
        custom_fields = application.custom_fields if isinstance(application.custom_fields, dict) else {}
        extra_keys.update(str(key) for key in custom_fields.keys() if str(key) not in known and str(key) not in SYSTEM_FIELD_IDS)

    for key in sorted(extra_keys):
        fields.append({"id": key, "label": DEFAULT_FIELD_LABELS.get(key, f"Доп. поле: {key}"), "type": "text", "required": False, "options": []})

    return fields


def _safe_filename(value: str, suffix: str) -> str:
    base = re.sub(r"[^0-9A-Za-zА-Яа-яЁё._ -]+", "", value or "event").strip(" ._- ")
    base = re.sub(r"\s+", "_", base)[:80] or "event"
    return f"{base}_{suffix}"


def _content_disposition(filename: str) -> str:
    fallback = re.sub(r"[^0-9A-Za-z_.-]+", "_", filename)
    return f"attachment; filename=\"{fallback}\"; filename*=UTF-8''{quote(filename)}"


def _build_response(content: bytes, content_type: str, filename: str) -> HttpResponse:
    response = HttpResponse(content, content_type=content_type)
    response["Content-Disposition"] = _content_disposition(filename)
    response["Cache-Control"] = "no-store"
    return response


def build_event_details_docx_response(event: Event) -> HttpResponse:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title = document.add_heading(event.name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("Основная информация", level=1)
    info_rows = [
        ("Название", event.name),
        ("Этап", event.stage),
        ("Дата начала", _format_date(event.start_date)),
        ("Дата завершения", _format_date(event.end_date)),
        ("Срок приема заявок", _format_datetime(event.end_app_date)),
        ("Организаторы", ", ".join(_display_user(user) for user in event.organizers.all()) or _display_user(event.leader)),
        ("Специализации", ", ".join(_event_specializations(event))),
        ("Ссылка на орг. чат VK", event.org_chat_url),
        ("ID беседы VK", event.org_chat_peer_id or ""),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in info_rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = _stringify(value)

    document.add_heading("Описание", level=1)
    description = (event.description or "").strip()
    if description:
        for paragraph in description.splitlines():
            document.add_paragraph(paragraph.strip() or " ")
    else:
        document.add_paragraph("Описание не заполнено.")

    directions = list(
        Direction.objects.filter(event=event)
        .select_related("leader")
        .prefetch_related("projects__curator")
        .order_by("name")
    )
    document.add_heading("Направления и проекты", level=1)
    if not directions:
        document.add_paragraph("Направления не добавлены.")
    for direction in directions:
        document.add_heading(direction.name, level=2)
        if direction.description:
            document.add_paragraph(direction.description)
        projects = list(Project.objects.filter(direction=direction).select_related("curator").order_by("name"))
        if not projects:
            document.add_paragraph("Проекты не добавлены.")
            continue
        projects_table = document.add_table(rows=1, cols=4)
        projects_table.style = "Table Grid"
        headers = ["Проект", "Куратор", "Команд", "Описание"]
        for index, header in enumerate(headers):
            projects_table.rows[0].cells[index].text = header
        for project in projects:
            cells = projects_table.add_row().cells
            cells[0].text = project.name
            cells[1].text = _display_user(project.curator)
            cells[2].text = _stringify(project.teams)
            cells[3].text = project.description or ""

    applications = list(Application.objects.filter(event=event).select_related("status"))
    status_counts = Counter(application.status.name if application.status_id else "Без статуса" for application in applications)
    document.add_heading("Заявки", level=1)
    document.add_paragraph(f"Всего заявок: {len(applications)}")
    if status_counts:
        status_table = document.add_table(rows=1, cols=2)
        status_table.style = "Table Grid"
        status_table.rows[0].cells[0].text = "Статус"
        status_table.rows[0].cells[1].text = "Количество"
        for status_name, count in status_counts.most_common():
            cells = status_table.add_row().cells
            cells[0].text = status_name
            cells[1].text = str(count)

    document.add_heading("Поля формы заявки", level=1)
    form_table = document.add_table(rows=1, cols=4)
    form_table.style = "Table Grid"
    for index, header in enumerate(["Поле", "Тип", "Обязательное", "Варианты"]):
        form_table.rows[0].cells[index].text = header
    for field in _normalize_form_fields(event.application_form_fields):
        cells = form_table.add_row().cells
        cells[0].text = _stringify(field["label"])
        cells[1].text = FIELD_TYPE_LABELS.get(str(field["type"]), str(field["type"]))
        cells[2].text = "Да" if field.get("required") else "Нет"
        cells[3].text = ", ".join(field.get("options") or [])

    stream = BytesIO()
    document.save(stream)
    filename = _safe_filename(event.name, "данные_мероприятия.docx")
    return _build_response(
        stream.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename,
    )


def build_event_applications_xlsx_response(event: Event) -> HttpResponse:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    applications = list(
        Application.objects.filter(event=event)
        .select_related("user", "user__crm_profile", "direction", "project", "specialization", "status")
        .order_by("date_sub", "id")
    )
    custom_fields = _custom_export_fields(event, applications)

    columns = [
        ("id", "ID заявки"),
        ("student", "ФИО"),
        ("email", "Email"),
        ("vk", "VK"),
        ("university", "Университет"),
        ("course", "Курс"),
        ("event", "Мероприятие"),
        ("direction", "Направление"),
        ("project", "Проект"),
        ("specialization", "Специализация"),
        ("status", "Текущий статус"),
        ("team", "ID команды"),
        ("tests", "Тесты назначены"),
        ("test_session", "Сессия тестирования"),
        ("created", "Дата подачи"),
        ("deadline", "Срок заявки"),
        ("comment", "Комментарий организатора"),
    ]
    columns.extend((f"custom:{field['id']}", str(field["label"])) for field in custom_fields)

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Заявки"

    header_fill = PatternFill("solid", fgColor="D9EAF7")
    header_font = Font(bold=True)
    border = Border(
        left=Side(style="thin", color="B7C9D8"),
        right=Side(style="thin", color="B7C9D8"),
        top=Side(style="thin", color="B7C9D8"),
        bottom=Side(style="thin", color="B7C9D8"),
    )

    for col_index, (_, title) in enumerate(columns, start=1):
        cell = sheet.cell(row=1, column=col_index, value=title)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    def cell_value(application: Application, key: str):
        user = application.user
        custom = application.custom_fields if isinstance(application.custom_fields, dict) else {}
        if key.startswith("custom:"):
            return custom.get(key.split(":", 1)[1], "")
        values = {
            "id": application.id,
            "student": _display_user(user),
            "email": getattr(user, "email", ""),
            "vk": _profile_value(user, "vk"),
            "university": _profile_value(user, "university"),
            "course": _profile_value(user, "course"),
            "event": event.name,
            "direction": application.direction.name if application.direction_id else "",
            "project": application.project.name if application.project_id else "",
            "specialization": application.specialization.name if application.specialization_id else "",
            "status": application.status.name if application.status_id else "",
            "team": application.team_id or "",
            "tests": application.tests_assigned,
            "test_session": application.test_session_id,
            "created": _format_datetime(application.date_sub),
            "deadline": _format_datetime(application.date_end),
            "comment": application.comment,
        }
        return values.get(key, "")

    for row_index, application in enumerate(applications, start=2):
        for col_index, (key, _) in enumerate(columns, start=1):
            cell = sheet.cell(row=row_index, column=col_index, value=_stringify(cell_value(application, key)))
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    for col_index, (_, title) in enumerate(columns, start=1):
        values = [str(sheet.cell(row=row, column=col_index).value or "") for row in range(1, min(sheet.max_row, 30) + 1)]
        width = min(max([len(title), *(len(value) for value in values)]) + 2, 42)
        sheet.column_dimensions[get_column_letter(col_index)].width = width

    stream = BytesIO()
    workbook.save(stream)
    filename = _safe_filename(event.name, "заявки.xlsx")
    return _build_response(
        stream.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename,
    )
