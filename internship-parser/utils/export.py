"""
Генерация файлов экспорта стажировок: CSV, Excel, Word.
"""
import csv
import io
from datetime import datetime
from typing import BinaryIO

from models.internship import Internship


FIELD_MAP = [
    "id",
    "title",
    "direction",
    "company",
    "city",
    "work_format",
    "link",
    "salary_from",
    "description",
]

FIELD_HEADERS = [
    "ID",
    "Название",
    "Направление",
    "Компания",
    "Город",
    "Формат работы",
    "Ссылка",
    "Зарплата от (руб.)",
    "Описание",
]


def _get_rows(session):
    internships = session.query(Internship).all()
    rows = []
    for i in internships:
        rows.append({
            "id": str(i.id),
            "title": i.title or "",
            "direction": i.direction or "",
            "company": i.company or "",
            "city": i.city or "",
            "work_format": i.work_format or "",
            "link": i.link or "",
            "salary_from": i.salary_from or "",
            "description": i.description or "",
        })
    return rows


def get_export_csv(rows) -> BinaryIO:
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=FIELD_MAP)
    writer.writeheader()
    for row in rows:
        writer.writerow(row)
    content = output.getvalue()
    output.close()
    return io.BytesIO(content.encode("utf-8-sig"))


def get_export_excel(rows) -> BinaryIO:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Стажировки"
    ws.append(FIELD_HEADERS)
    for row in rows:
        ws.append([row[field] for field in FIELD_MAP])
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def get_export_word(rows) -> BinaryIO:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    doc.add_heading("Стажировки", level=1)
    doc.add_paragraph(datetime.now().strftime("Сформировано: %d.%m.%Y %H:%M"))

    if rows:
        table = doc.add_table(rows=1, cols=len(FIELD_HEADERS))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, header in enumerate(FIELD_HEADERS):
            table.rows[0].cells[i].text = header
        for row in rows:
            cells = table.add_row().cells
            for i, field in enumerate(FIELD_MAP):
                cells[i].text = str(row[field]) if row[field] is not None else ""
    else:
        doc.add_paragraph("Нет данных")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
